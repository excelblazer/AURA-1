import os
from docx import Document
from docx.shared import Pt, Inches
import datetime
from typing import List, Dict, Any
import pandas as pd

def generate_progress_reports(month: str, year: int, feedback_data: Dict = None) -> List[str]:
    """
    Generate Progress Reports for all students for the specified month and year
    
    Args:
        month (str): Month name
        year (int): Year
        feedback_data (Dict): Data extracted from the Daily Feedback Sheet
        
    Returns:
        List[str]: List of generated file paths
    """
    # In a real app, this data would come from the database
    # For demonstration, we'll use a simplified example
    
    if feedback_data is None:
        # Mock data for demonstration
        feedback_data = {
            'students': [
                {
                    'id': 'S001',
                    'first_name': 'John',
                    'last_name': 'Doe',
                    'grade': '8',
                    'case_number': 'CN12345',
                    'tutor_first_name': 'Jane',
                    'tutor_last_name': 'Smith',
                    'caregiver_name': 'Robert Doe',
                    'caregiver_phone': '555-123-4567',
                    'sessions': [
                        {
                            'date': datetime.datetime(year, {'January': 1, 'February': 2, 'March': 3, 'April': 4, 'May': 5, 'June': 6, 'July': 7, 'August': 8, 'September': 9, 'October': 10, 'November': 11, 'December': 12}.get(month, 1), 5),
                            'goal': 'Improve math skills',
                            'feedback': 'Good progress in equations',
                            'areas_of_focus': 'Algebra, Fractions',
                            'achievements': 'Completed all assigned problems',
                            'challenges': 'Struggles with word problems',
                            'next_steps': 'Focus on word problem strategies'
                        },
                        {
                            'date': datetime.datetime(year, {'January': 1, 'February': 2, 'March': 3, 'April': 4, 'May': 5, 'June': 6, 'July': 7, 'August': 8, 'September': 9, 'October': 10, 'November': 11, 'December': 12}.get(month, 1), 12),
                            'goal': 'Reading comprehension',
                            'feedback': 'Improved understanding of passages',
                            'areas_of_focus': 'Main idea, Context clues',
                            'achievements': 'Increased reading speed',
                            'challenges': 'Vocabulary retention',
                            'next_steps': 'Vocabulary building exercises'
                        }
                    ]
                },
                {
                    'id': 'S002',
                    'first_name': 'Emma',
                    'last_name': 'Johnson',
                    'grade': '6',
                    'case_number': 'CN67890',
                    'tutor_first_name': 'Michael',
                    'tutor_last_name': 'Brown',
                    'caregiver_name': 'Sarah Johnson',
                    'caregiver_phone': '555-987-6543',
                    'sessions': [
                        {
                            'date': datetime.datetime(year, {'January': 1, 'February': 2, 'March': 3, 'April': 4, 'May': 5, 'June': 6, 'July': 7, 'August': 8, 'September': 9, 'October': 10, 'November': 11, 'December': 12}.get(month, 1), 8),
                            'goal': 'Improve writing skills',
                            'feedback': 'Good progress in essay structure',
                            'areas_of_focus': 'Paragraph structure, Transitions',
                            'achievements': 'Completed 2-page essay',
                            'challenges': 'Spelling and grammar',
                            'next_steps': 'Practice using grammar checker'
                        }
                    ]
                }
            ]
        }
    
    # Create output directory if it doesn't exist
    os.makedirs("outputs", exist_ok=True)
    
    # List to store generated file paths
    generated_files = []
    
    # Month as a number for file naming
    month_num = {'January': '01', 'February': '02', 'March': '03', 'April': '04', 
                 'May': '05', 'June': '06', 'July': '07', 'August': '08', 
                 'September': '09', 'October': '10', 'November': '11', 'December': '12'}.get(month, '01')
    
    # Generate a report for each student
    for student in feedback_data['students']:
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Client1 Progress Report', 0)
        title.alignment = 1  # Center alignment
        
        # Add reporting period
        doc.add_paragraph(f"Reporting Period: {month} {year}")
        
        # Add student information section
        doc.add_heading('Student Information', level=1)
        student_info = doc.add_paragraph()
        student_info.add_run('Student Name: ').bold = True
        student_info.add_run(f"{student['first_name']} {student['last_name']}")
        student_info.add_run('\nGrade: ').bold = True
        student_info.add_run(f"{student['grade']}")
        student_info.add_run('\nCase Number: ').bold = True
        student_info.add_run(f"{student['case_number']}")
        
        # Add tutor information
        doc.add_heading('Tutor Information', level=1)
        tutor_info = doc.add_paragraph()
        tutor_info.add_run('Tutor Name: ').bold = True
        tutor_info.add_run(f"{student['tutor_first_name']} {student['tutor_last_name']}")
        
        # Add progress summary
        doc.add_heading('Monthly Progress Summary', level=1)
        
        # Collect all feedback from sessions
        areas_of_focus = set()
        achievements = []
        challenges = []
        next_steps = []
        
        for session in student['sessions']:
            if 'areas_of_focus' in session:
                for area in session['areas_of_focus'].split(','):
                    areas_of_focus.add(area.strip())
            if 'achievements' in session:
                achievements.append(session['achievements'])
            if 'challenges' in session:
                challenges.append(session['challenges'])
            if 'next_steps' in session:
                next_steps.append(session['next_steps'])
        
        # Areas of Focus
        doc.add_heading('Areas of Focus', level=2)
        focus_para = doc.add_paragraph()
        if areas_of_focus:
            focus_para.add_run(', '.join(areas_of_focus))
        else:
            focus_para.add_run('No specific areas of focus recorded for this month.')
        
        # Achievements
        doc.add_heading('Achievements', level=2)
        if achievements:
            for achievement in achievements:
                doc.add_paragraph(f"• {achievement}", style='List Bullet')
        else:
            doc.add_paragraph('No specific achievements recorded for this month.')
        
        # Challenges
        doc.add_heading('Challenges', level=2)
        if challenges:
            for challenge in challenges:
                doc.add_paragraph(f"• {challenge}", style='List Bullet')
        else:
            doc.add_paragraph('No specific challenges recorded for this month.')
        
        # Next Steps
        doc.add_heading('Next Steps', level=2)
        if next_steps:
            for step in next_steps:
                doc.add_paragraph(f"• {step}", style='List Bullet')
        else:
            doc.add_paragraph('No specific next steps recorded for this month.')
        
        # Add signature section
        doc.add_heading('Signatures', level=1)
        
        signatures = doc.add_paragraph()
        signatures.add_run('Tutor Signature: ').bold = True
        signatures.add_run('_______________________     ')
        signatures.add_run('Date: ').bold = True
        signatures.add_run('_________________\n\n')
        
        signatures.add_run('Agency Representative: ').bold = True
        signatures.add_run('_______________________     ')
        signatures.add_run('Date: ').bold = True
        signatures.add_run('_________________')
        
        # Save the document
        filename = f"PR_{student['last_name']}_{student['first_name']}_{month}_{year}.docx"
        filepath = os.path.join("outputs", filename)
        doc.save(filepath)
        
        # Convert to PDF (in a real app)
        # For this example, we'll just add the docx file
        generated_files.append(filepath)
        
        # In a real app, you would convert to PDF here
        # pdf_filepath = os.path.join("outputs", f"PR_{student['last_name']}_{student['first_name']}_{month}_{year}.pdf")
        # convert_to_pdf(filepath, pdf_filepath)
        # generated_files.append(pdf_filepath)
    
    return generated_files

def convert_to_pdf(docx_path, pdf_path):
    """
    Convert a Word document to PDF
    In a real application, you would use a library like docx2pdf or an external service
    
    For this example, we'll just pass
    """
    # Placeholder for PDF conversion functionality
    pass