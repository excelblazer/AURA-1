from typing import Dict, List, Any, Optional
import os
from datetime import datetime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter
from ..database import crud

def generate_attendance_record(
    student: Dict[str, Any], 
    sessions: List[Dict[str, Any]], 
    month: int, 
    year: int,
    template_id: Optional[str] = None
) -> str:
    """
    Generate Attendance Record for a student
    
    Args:
        student: Dictionary containing student data
        sessions: List of sessions for the student
        month: Month for the report
        year: Year for the report
        template_id: Optional template ID to use
        
    Returns:
        Path to the generated file
    """
    # Create output directory if it doesn't exist
    output_dir = os.path.join("output", "attendance_records")
    os.makedirs(output_dir, exist_ok=True)
    
    # Create document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = f"Attendance Record - {student.get('first_name')} {student.get('last_name')}"
    doc.core_properties.author = "Client1"
    
    # Add title
    title = doc.add_heading(f"ATTENDANCE RECORD - {datetime.strptime(str(month), '%m').strftime('%B')} {year}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student information section
    doc.add_paragraph().add_run("STUDENT INFORMATION").bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Fill student information
    cells = table.rows[0].cells
    cells[0].text = "Student Name:"
    cells[1].text = f"{student.get('first_name')} {student.get('last_name')}"
    
    cells = table.rows[1].cells
    cells[0].text = "Grade:"
    cells[1].text = str(student.get('grade', ''))
    
    cells = table.rows[2].cells
    cells[0].text = "Case Number:"
    cells[1].text = student.get('case_number', '')
    
    doc.add_paragraph()
    
    # Add tutor information section
    doc.add_paragraph().add_run("TUTOR INFORMATION").bold = True
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Fill tutor information
    cells = table.rows[0].cells
    cells[0].text = "Tutor Name:"
    cells[1].text = student.get('tutor_assigned', '')
    
    cells = table.rows[1].cells
    cells[0].text = "Tutoring Start Date:"
    cells[1].text = student.get('start_date', '')
    
    doc.add_paragraph()
    
    # Add caregiver information section
    doc.add_paragraph().add_run("CAREGIVER INFORMATION").bold = True
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Fill caregiver information
    cells = table.rows[0].cells
    cells[0].text = "Caregiver Name:"
    cells[1].text = student.get('caregiver_name', '')
    
    cells = table.rows[1].cells
    cells[0].text = "Caregiver Phone:"
    cells[1].text = student.get('caregiver_phone', '')
    
    doc.add_paragraph()
    
    # Calculate total hours
    total_hours = sum(session.get('hours', 0) for session in sessions)
    
    # Add total hours
    p = doc.add_paragraph()
    p.add_run(f"Total Hours: {total_hours:.2f}").bold = True
    
    doc.add_paragraph()
    
    # Add sessions table
    doc.add_paragraph().add_run("SESSION DETAILS").bold = True
    
    # Create table header
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = "Date"
    header_cells[1].text = "Start Time"
    header_cells[2].text = "End Time"
    header_cells[3].text = "Hours"
    header_cells[4].text = "Goal/Objective"
    
    # Add session data
    for session in sessions:
        row_cells = table.add_row().cells
        row_cells[0].text = session.get('date', '')
        row_cells[1].text = session.get('start_time', '')
        row_cells[2].text = session.get('end_time', '')
        row_cells[3].text = f"{session.get('hours', 0):.2f}"
        row_cells[4].text = session.get('goal', '')
    
    doc.add_paragraph()
    
    # Add signature lines
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Caregiver Signature")
    
    doc.add_paragraph()
    
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Tutor Signature")
    
    doc.add_paragraph()
    
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Agency Representative Signature")
    
    # Save document
    filename = f"{student.get('last_name')}_{student.get('first_name')}_AR_{month}_{year}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    
    return file_path

def generate_progress_report(
    student: Dict[str, Any], 
    sessions: List[Dict[str, Any]], 
    month: int, 
    year: int,
    template_id: Optional[str] = None
) -> str:
    """
    Generate Progress Report for a student
    
    Args:
        student: Dictionary containing student data
        sessions: List of sessions for the student
        month: Month for the report
        year: Year for the report
        template_id: Optional template ID to use
        
    Returns:
        Path to the generated file
    """
    # Create output directory if it doesn't exist
    output_dir = os.path.join("output", "progress_reports")
    os.makedirs(output_dir, exist_ok=True)
    
    # Create document
    doc = docx.Document()
    
    # Set document properties
    doc.core_properties.title = f"Progress Report - {student.get('first_name')} {student.get('last_name')}"
    doc.core_properties.author = "Client1"
    
    # Add title
    title = doc.add_heading(f"PROGRESS REPORT - {datetime.strptime(str(month), '%m').strftime('%B')} {year}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student information section
    doc.add_paragraph().add_run("STUDENT INFORMATION").bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Fill student information
    cells = table.rows[0].cells
    cells[0].text = "Student Name:"
    cells[1].text = f"{student.get('first_name')} {student.get('last_name')}"
    
    cells = table.rows[1].cells
    cells[0].text = "Grade:"
    cells[1].text = str(student.get('grade', ''))
    
    cells = table.rows[2].cells
    cells[0].text = "Case Number:"
    cells[1].text = student.get('case_number', '')
    
    doc.add_paragraph()
    
    # Add tutor information
    doc.add_paragraph().add_run("TUTOR INFORMATION").bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Fill tutor information
    cells = table.rows[0].cells
    cells[0].text = "Tutor Name:"
    cells[1].text = student.get('tutor_assigned', '')
    
    doc.add_paragraph()
    
    # Add progress summary
    doc.add_paragraph().add_run("PROGRESS SUMMARY").bold = True
    
    # Collect all feedback from sessions
    all_feedback = []
    for session in sessions:
        if session.get('feedback'):
            all_feedback.append(session.get('feedback'))
    
    # Add feedback paragraphs
    if all_feedback:
        for feedback in all_feedback:
            doc.add_paragraph(feedback)
    else:
        doc.add_paragraph("No feedback provided for this period.")
    
    doc.add_paragraph()
    
    # Add goals section
    doc.add_paragraph().add_run("GOALS AND OBJECTIVES").bold = True
    
    # Collect all goals from sessions
    all_goals = []
    for session in sessions:
        if session.get('goal') and session.get('goal') not in all_goals:
            all_goals.append(session.get('goal'))
    
    # Add goals as bullet points
    if all_goals:
        for goal in all_goals:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(goal)
    else:
        doc.add_paragraph("No goals specified for this period.")
    
    doc.add_paragraph()
    
    # Add recommendations section
    doc.add_paragraph().add_run("RECOMMENDATIONS").bold = True
    doc.add_paragraph("Based on the student's progress this month, the following recommendations are made:")
    
    # Add placeholder recommendations
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Continue with current tutoring plan")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Focus on areas needing improvement")
    
    doc.add_paragraph()
    
    # Add signature lines
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Tutor Signature")
    
    doc.add_paragraph()
    
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Agency Representative Signature")
    
    # Save document
    filename = f"{student.get('last_name')}_{student.get('first_name')}_PR_{month}_{year}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    
    return file_path

def generate_invoice(
    students: List[Dict[str, Any]], 
    sessions: List[Dict[str, Any]], 
    month: int, 
    year: int,
    template_id: Optional[str] = None
) -> str:
    """
    Generate Invoice for all students
    
    Args:
        students: List of students
        sessions: List of all sessions
        month: Month for the invoice
        year: Year for the invoice
        template_id: Optional template ID to use
        
    Returns:
        Path to the generated file
    """
    # Create output directory if it doesn't exist
    output_dir = os.path.join("output", "invoices")
    os.makedirs(output_dir, exist_ok=True)
    
    # Create workbook
    filename = f"Client1_Invoice_{month}_{year}.xlsx"
    file_path = os.path.join(output_dir, filename)
    
    workbook = xlsxwriter.Workbook(file_path)
    worksheet = workbook.add_worksheet("Invoice")
    
    # Add header
    bold_format = workbook.add_format({'bold': True})
    header_format = workbook.add_format({'bold': True, 'align': 'center', 'bg_color': '#D3D3D3', 'border': 1})
    
    worksheet.write(0, 0, "Client1 Invoice", bold_format)
    worksheet.write(1, 0, f"Month: {datetime.strptime(str(month), '%m').strftime('%B')}", bold_format)
    worksheet.write(1, 1, f"Year: {year}", bold_format)
    
    # Add table headers
    headers = ["Student Name", "Grade", "Case Number", "Tutor", "Hours", "Rate", "Total"]
    for col, header in enumerate(headers):
        worksheet.write(3, col, header, header_format)
    
    # Set column widths
    worksheet.set_column(0, 0, 25)  # Student Name
    worksheet.set_column(1, 1, 10)  # Grade
    worksheet.set_column(2, 2, 15)  # Case Number
    worksheet.set_column(3, 3, 25)  # Tutor
    worksheet.set_column(4, 4, 10)  # Hours
    worksheet.set_column(5, 5, 10)  # Rate
    worksheet.set_column(6, 6, 15)  # Total
    
    # Add student data
    row = 4
    total_invoice_amount = 0
    
    # Format for currency
    currency_format = workbook.add_format({'num_format': '$#,##0.00'})
    
    for student in students:
        student_name = f"{student.get('first_name')} {student.get('last_name')}"
        
        # Calculate total hours for student
        student_sessions = [s for s in sessions if s.get('student_name') == student_name]
        total_hours = sum(session.get('hours', 0) for session in student_sessions)
        
        # Use default rate of $50 per hour
        rate = 50
        total_amount = total_hours * rate
        total_invoice_amount += total_amount
        
        # Write student data
        worksheet.write(row, 0, student_name)
        worksheet.write(row, 1, student.get('grade', ''))
        worksheet.write(row, 2, student.get('case_number', ''))
        worksheet.write(row, 3, student.get('tutor_assigned', ''))
        worksheet.write(row, 4, total_hours)
        worksheet.write(row, 5, rate, currency_format)
        worksheet.write(row, 6, total_amount, currency_format)
        
        row += 1
    
    # Add total row
    worksheet.write(row + 1, 5, "TOTAL:", bold_format)
    worksheet.write(row + 1, 6, total_invoice_amount, currency_format)
    
    # Add signature line
    worksheet.write(row + 4, 0, "Authorized Signature:")
    worksheet.write(row + 4, 3, "Date:")
    
    workbook.close()
    
    return file_path

def generate_service_log(
    students: List[Dict[str, Any]], 
    sessions: List[Dict[str, Any]], 
    month: int, 
    year: int,
    template_id: Optional[str] = None
) -> str:
    """
    Generate Agency Service Log
    
    Args:
        students: List of students
        sessions: List of all sessions
        month: Month for the service log
        year: Year for the service log
        template_id: Optional template ID to use
        
    Returns:
        Path to the generated file
    """
    # Create output directory if it doesn't exist
    output_dir = os.path.join("output", "service_logs")
    os.makedirs(output_dir, exist_ok=True)
    
    # Create workbook
    filename = f"Client1_Service_Log_{month}_{year}.xlsx"
    file_path = os.path.join(output_dir, filename)
    
    workbook = xlsxwriter.Workbook(file_path)
    worksheet = workbook.add_worksheet("Service Log")
    
    # Add header
    bold_format = workbook.add_format({'bold': True})
    header_format = workbook.add_format({'bold': True, 'align': 'center', 'bg_color': '#D3D3D3', 'border': 1})
    date_format = workbook.add_format({'num_format': 'mm/dd/yyyy'})
    
    worksheet.write(0, 0, "Client1 Agency Service Log", bold_format)
    worksheet.write(1, 0, f"Month: {datetime.strptime(str(month), '%m').strftime('%B')}", bold_format)
    worksheet.write(1, 1, f"Year: {year}", bold_format)
    
    # Add table headers
    headers = ["Date", "Student Name", "Case Number", "Tutor", "Service Type", "Start Time", "End Time", "Hours", "Notes"]
    for col, header in enumerate(headers):
        worksheet.write(3, col, header, header_format)
    
    # Set column widths
    worksheet.set_column(0, 0, 12)  # Date
    worksheet.set_column(1, 1, 25)  # Student Name
    worksheet.set_column(2, 2, 15)  # Case Number
    worksheet.set_column(3, 3, 25)  # Tutor
    worksheet.set_column(4, 4, 15)  # Service Type
    worksheet.set_column(5, 5, 12)  # Start Time
    worksheet.set_column(6, 6, 12)  # End Time
    worksheet.set_column(7, 7, 10)  # Hours
    worksheet.set_column(8, 8, 40)  # Notes
    
    # Add session data
    row = 4
    
    for session in sessions:
        student_name = session.get('student_name', '')
        
        # Find student data
        student = next((s for s in students if f"{s.get('first_name')} {s.get('last_name')}" == student_name), None)
        
        if student:
            # Write session data
            worksheet.write(row, 0, session.get('date', ''))
            worksheet.write(row, 1, student_name)
            worksheet.write(row, 2, student.get('case_number', ''))
            worksheet.write(row, 3, student.get('tutor_assigned', ''))
            worksheet.write(row, 4, "Tutoring")
            worksheet.write(row, 5, session.get('start_time', ''))
            worksheet.write(row, 6, session.get('end_time', ''))
            worksheet.write(row, 7, session.get('hours', 0))
            worksheet.write(row, 8, session.get('goal', ''))
            
            row += 1
    
    # Add signature line
    worksheet.write(row + 2, 0, "Agency Representative Signature:")
    worksheet.write(row + 2, 4, "Date:")
    
    workbook.close()
    
    return file_path
