import os
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_attendance_records(month: str, year: int, data: Dict[str, Any] = None) -> List[str]:
    """
    Generate attendance records for all students
    
    Args:
        month: Month name (e.g., "January")
        year: Year (e.g., 2023)
        data: Extracted data from feedback sheets (if None, uses mock data)
        
    Returns:
        List of file paths to generated attendance records
    """
    # Create output directory if it doesn't exist
    output_dir = os.path.join("outputs", f"{month}_{year}")
    os.makedirs(output_dir, exist_ok=True)
    
    # Use mock data if none provided
    if data is None:
        data = generate_mock_data()
    
    generated_files = []
    
    # Generate AR for each student
    for student in data.get("students", []):
        # Skip students with no tutoring sessions
        student_sessions = [s for s in data.get("sessions", []) if s.get("student_id") == student.get("id")]
        if not student_sessions:
            continue
        
        # Generate the attendance record
        file_path = generate_student_ar(student, student_sessions, month, year, output_dir)
        generated_files.append(file_path)
    
    return generated_files

def generate_student_ar(student: Dict[str, Any], sessions: List[Dict[str, Any]], 
                        month: str, year: int, output_dir: str) -> str:
    """Generate attendance record for a single student"""
    # Create a new document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("LOCOE GAIN ATTENDANCE RECORD")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add reporting period
    period = doc.add_paragraph()
    period_run = period.add_run(f"Reporting Month: {month} {year}")
    period_run.bold = True
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student information section
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("Student Name: ").bold = True
    info.add_run(f"{student.get('last_name')}, {student.get('first_name')}")
    
    info = doc.add_paragraph()
    info.add_run("Grade: ").bold = True
    info.add_run(f"{student.get('grade', '')}")
    
    info = doc.add_paragraph()
    info.add_run("Case Number: ").bold = True
    info.add_run(f"{student.get('case_number', '')}")
    
    info = doc.add_paragraph()
    info.add_run("Tutoring Start Date: ").bold = True
    info.add_run(f"{student.get('tutor_start_date', '')}")
    
    # Add tutor information
    doc.add_paragraph()
    tutor_section = doc.add_paragraph()
    tutor_section.add_run("Tutor Name: ").bold = True
    tutor_section.add_run(f"{student.get('tutor_assigned', '')}")
    
    # Add caregiver information
    caregiver_section = doc.add_paragraph()
    caregiver_section.add_run("Caregiver Name: ").bold = True
    caregiver_section.add_run(f"{student.get('caregiver_name', '')}")
    
    caregiver_phone = doc.add_paragraph()
    caregiver_phone.add_run("Caregiver Phone: ").bold = True
    caregiver_phone.add_run(f"{student.get('caregiver_phone', '')}")
    
    # Calculate total hours
    total_hours = sum(session.get("hours", 0) for session in sessions)
    
    total_section = doc.add_paragraph()
    total_section.add_run("Total Hours: ").bold = True
    total_section.add_run(f"{total_hours:.2f}")
    
    # Add session table
    doc.add_paragraph()
    table_header = doc.add_paragraph()
    table_header.add_run("Session Details:").bold = True
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Date"
    header_cells[1].text = "Start Time"
    header_cells[2].text = "End Time"
    header_cells[3].text = "Hours"
    header_cells[4].text = "Goals/Activities"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for session in sorted(sessions, key=lambda x: x.get("date", "")):
        row_cells = table.add_row().cells
        row_cells[0].text = session.get("date", "")
        row_cells[1].text = session.get("time_in", "")
        row_cells[2].text = session.get("time_out", "")
        row_cells[3].text = f"{session.get('hours', 0):.2f}"
        row_cells[4].text = session.get("goal", "")
    
    # Add signature section
    doc.add_paragraph()
    doc.add_paragraph()
    
    signatures = doc.add_paragraph()
    signatures.add_run("Signatures:").bold = True
    
    # Caregiver signature
    doc.add_paragraph()
    caregiver_sig = doc.add_paragraph()
    caregiver_sig.add_run("Caregiver Signature: ").bold = True
    caregiver_sig.add_run("_________________________________")
    caregiver_sig.add_run("    Date: ").bold = True
    caregiver_sig.add_run("_______________")
    
    # Tutor signature
    doc.add_paragraph()
    tutor_sig = doc.add_paragraph()
    tutor_sig.add_run("Tutor Signature: ").bold = True
    tutor_sig.add_run("_________________________________")
    tutor_sig.add_run("    Date: ").bold = True
    tutor_sig.add_run("_______________")
    
    # Agency signature
    doc.add_paragraph()
    agency_sig = doc.add_paragraph()
    agency_sig.add_run("Agency Representative: ").bold = True
    agency_sig.add_run("_________________________________")
    agency_sig.add_run("    Date: ").bold = True
    agency_sig.add_run("_______________")
    
    # Save the document
    file_name = f"AR_{student.get('last_name')}_{student.get('first_name')}_{month}_{year}.docx"
    file_path = os.path.join(output_dir, file_name)
    doc.save(file_path)
    
    # In a real application, you'd convert to PDF here
    pdf_path = file_path.replace(".docx", ".pdf")
    convert_to_pdf(file_path, pdf_path)
    
    return pdf_path

def convert_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert DOCX to PDF
    
    In a real application, you would use a library like docx2pdf or a service like LibreOffice.
    For this example, we'll just simulate the conversion.
    """
    # Simulate conversion by creating an empty file
    with open(pdf_path, "w") as f:
        f.write("PDF content would be here in a real application")
    
    return pdf_path

def generate_mock_data() -> Dict[str, Any]:
    """Generate mock data for testing"""
    return {
        "students": [
            {
                "id": "student1",
                "first_name": "John",
                "last_name": "Doe",
                "full_name": "John Doe",
                "grade": "9",
                "case_number": "CASE123",
                "tutor_start_date": "01/15/2023",
                "tutor_assigned": "Jane Smith",
                "caregiver_name": "Mary Doe",
                "caregiver_phone": "555-123-4567",
                "caregiver_email": "mary.doe@example.com"
            }
        ],
        "sessions": [
            {
                "student_id": "student1",
                "student_name": "John Doe",
                "date": "03/05/2023",
                "time_in": "3:00 PM",
                "time_out": "5:00 PM",
                "hours": 2.0,
                "goal": "Math homework and algebra review",
                "is_no_show": False
            },
            {
                "student_id": "student1",
                "student_name": "John Doe",
                "date": "03/12/2023",
                "time_in": "3:00 PM",
                "time_out": "5:00 PM",
                "hours": 2.0,
                "goal": "Science project research",
                "is_no_show": False
            },
            {
                "student_id": "student1",
                "student_name": "John Doe",
                "date": "03/19/2023",
                "time_in": "3:00 PM",
                "time_out": "5:00 PM",
                "hours": 2.0,
                "goal": "English essay writing",
                "is_no_show": False
            },
            {
                "student_id": "student1",
                "student_name": "John Doe",
                "date": "03/26/2023",
                "time_in": "3:00 PM",
                "time_out": "5:00 PM",
                "hours": 2.0,
                "goal": "Test preparation",
                "is_no_show": False
            }
        ]
    }